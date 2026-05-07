"""
Tests for DOCX text extraction processor.
"""
from __future__ import annotations

import io

import pytest

from processors.docx_processor import DocxProcessor


def _make_docx_bytes(paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> bytes:
    """Helper: create a .docx file in memory with given paragraphs and optional tables."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)

    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = len(table_data[0]) if rows > 0 else 0
            if rows and cols:
                table = doc.add_table(rows=rows, cols=cols)
                for r, row_data in enumerate(table_data):
                    for c, cell_text in enumerate(row_data):
                        table.rows[r].cells[c].text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_docx_extract_simple(tmp_path):
    """Extract text from a simple DOCX with paragraphs."""
    proc = DocxProcessor(upload_dir=str(tmp_path))
    content = _make_docx_bytes([
        "Заголовок документа",
        "Первый абзац с текстом.",
        "Второй абзац документа.",
    ])
    result = await proc.extract_text("test.docx", content)
    assert result.filename == "test.docx"
    assert "Заголовок документа" in result.text
    assert "Первый абзац" in result.text
    assert "Второй абзац" in result.text
    assert result.page_count >= 1


@pytest.mark.asyncio
async def test_docx_extract_empty(tmp_path):
    """Extract from a DOCX with no paragraphs."""
    proc = DocxProcessor(upload_dir=str(tmp_path))
    content = _make_docx_bytes([])
    result = await proc.extract_text("empty.docx", content)
    assert result.filename == "empty.docx"
    assert result.text == ""
    assert result.page_count >= 1


@pytest.mark.asyncio
async def test_docx_extract_with_tables(tmp_path):
    """Extract text from a DOCX containing a table."""
    proc = DocxProcessor(upload_dir=str(tmp_path))
    content = _make_docx_bytes(
        paragraphs=["Invoice Header"],
        tables=[
            [
                ["Item", "Qty", "Price"],
                ["Widget", "10", "100"],
                ["Gadget", "5", "250"],
            ]
        ],
    )
    result = await proc.extract_text("invoice.docx", content)
    assert "Invoice Header" in result.text
    assert "Widget" in result.text
    assert "Gadget" in result.text
    # Table rows should be pipe-separated
    assert "|" in result.text


@pytest.mark.asyncio
async def test_docx_extract_whitespace_only(tmp_path):
    """DOCX with only whitespace paragraphs should yield empty text."""
    proc = DocxProcessor(upload_dir=str(tmp_path))
    content = _make_docx_bytes(["   ", "  ", ""])
    result = await proc.extract_text("whitespace.docx", content)
    assert result.filename == "whitespace.docx"
    # Whitespace-only paragraphs are filtered out
    assert result.text.strip() == ""


@pytest.mark.asyncio
async def test_docx_saves_file(tmp_path):
    """Verify the processor saves the file to upload_dir."""
    proc = DocxProcessor(upload_dir=str(tmp_path))
    content = _make_docx_bytes(["Hello world"])
    await proc.extract_text("saved.docx", content)
    assert (tmp_path / "saved.docx").exists()


@pytest.mark.asyncio
async def test_docx_cyrillic_text(tmp_path):
    """Ensure Cyrillic text is preserved correctly."""
    proc = DocxProcessor(upload_dir=str(tmp_path))
    content = _make_docx_bytes([
        "Договор №123 от 01.01.2024",
        "Исполнитель: ООО Рога и Копыта",
        "Сумма: 100 000 руб.",
    ])
    result = await proc.extract_text("contract.docx", content)
    assert "Договор" in result.text
    assert "Рога и Копыта" in result.text
    assert "100 000" in result.text


@pytest.mark.asyncio
async def test_docx_processor_init(tmp_path):
    """Processor creates upload dir on init."""
    subdir = tmp_path / "docx_uploads"
    _proc = DocxProcessor(upload_dir=str(subdir))
    assert subdir.exists()
