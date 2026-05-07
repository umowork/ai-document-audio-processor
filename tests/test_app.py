"""
Tests for Streamlit app module — import and function tests.
We don't run the Streamlit server; we test the helper functions directly.
"""
from __future__ import annotations

import sys
from unittest.mock import MagicMock

import pytest


def test_app_imports():
    """app.py should be importable without errors."""
    # We need to mock streamlit since it requires a running server
    mock_st = MagicMock()
    mock_st.set_page_config = MagicMock()
    mock_st.title = MagicMock()
    mock_st.caption = MagicMock()
    mock_st.sidebar = MagicMock()
    mock_st.tabs = MagicMock(return_value=[MagicMock(), MagicMock()])
    mock_st.file_uploader = MagicMock(return_value=None)
    mock_st.checkbox = MagicMock(return_value=False)
    mock_st.button = MagicMock(return_value=False)
    mock_st.info = MagicMock()
    mock_st.divider = MagicMock()
    mock_st.subheader = MagicMock()
    mock_st.header = MagicMock()
    mock_st.markdown = MagicMock()

    sys.modules["streamlit"] = mock_st

    try:
        # Import the app module
        import app

        # Verify key functions exist
        assert hasattr(app, "main")
        assert hasattr(app, "process_document_file")
        assert hasattr(app, "process_audio_file")
        assert hasattr(app, "run_async")
        assert hasattr(app, "get_config")
    finally:
        del sys.modules["streamlit"]
        if "app" in sys.modules:
            del sys.modules["app"]


@pytest.mark.asyncio
async def test_process_document_file_docx():
    """Test document processing with a DOCX file in mock mode."""
    import os
    os.environ["MOCK_MODE"] = "true"

    mock_st = MagicMock()
    sys.modules["streamlit"] = mock_st

    try:
        import io

        from docx import Document as DocxDocument

        from config import Config
        from processors.docx_processor import DocxProcessor
        from processors.structured_extractor import StructuredExtractor

        config = Config.from_env()
        docx_proc = DocxProcessor(upload_dir=config.upload_dir)
        extractor = StructuredExtractor(
            openai_api_key=config.openai_api_key,
            gigachat_api_key=config.gigachat_api_key,
            mock_mode=config.mock_mode,
        )

        # Create a simple DOCX
        doc = DocxDocument()
        doc.add_paragraph("Invoice #12345")
        doc.add_paragraph("Amount: $500")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = await docx_proc.extract_text("invoice.docx", buf.getvalue())
        schema = await extractor.extract(result.text)
        result.fields = schema.fields
        assert result.filename == "invoice.docx"
        assert "Invoice" in result.text
        assert len(result.fields) > 0  # mock mode produces fields
    finally:
        del sys.modules["streamlit"]
        if "app" in sys.modules:
            del sys.modules["app"]


@pytest.mark.asyncio
async def test_process_audio_file_mock():
    """Test audio processing in mock mode."""
    import os
    os.environ["MOCK_MODE"] = "true"

    mock_st = MagicMock()
    sys.modules["streamlit"] = mock_st

    try:
        from config import Config
        from processors.audio_whisper import AudioProcessor

        config = Config.from_env()
        audio_proc = AudioProcessor(
            openai_api_key=config.openai_api_key,
            upload_dir=config.upload_dir,
            mock_mode=config.mock_mode,
        )
        result = await audio_proc.transcribe("test.mp3", b"fake audio data")
        assert result.filename == "test.mp3"
        assert result.duration > 0
        assert len(result.segments) > 0
    finally:
        del sys.modules["streamlit"]
        if "app" in sys.modules:
            del sys.modules["app"]


def test_run_async():
    """Test that run_async executes a coroutine."""
    mock_st = MagicMock()
    sys.modules["streamlit"] = mock_st

    try:
        import app

        async def sample():
            return 42

        result = app.run_async(sample())
        assert result == 42
    finally:
        del sys.modules["streamlit"]
        if "app" in sys.modules:
            del sys.modules["app"]
