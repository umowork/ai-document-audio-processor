"""
Smoke / integration tests for project 03.
"""
from __future__ import annotations

import io

from fastapi.testclient import TestClient

from main import app

client = TestClient(app)


def test_health():
    resp = client.get("/health")
    assert resp.status_code == 200
    data = resp.json()
    assert data["status"] == "ok"


def test_upload_no_file():
    resp = client.post("/upload")
    assert resp.status_code == 422  # validation error


def test_upload_empty_file():
    resp = client.post(
        "/upload",
        files={"file": ("empty.txt", b"", "text/plain")},
        data={"process_type": "auto"},
    )
    assert resp.status_code == 400


def test_status_not_found():
    resp = client.get("/status/nonexistent-job-id")
    assert resp.status_code == 404


def test_download_not_ready():
    from tasks.process_document import job_manager
    job_manager.create_job(
        job_id="test-not-ready",
        filename="test.txt",
        content=b"hello",
    )
    resp = client.get("/download/test-not-ready")
    assert resp.status_code == 400


def test_process_docx_endpoint():
    """Upload a .docx through /process/document endpoint."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Test Document Title")
    doc.add_paragraph("Body paragraph with content.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    resp = client.post(
        "/process/document",
        files={
            "file": (
                "test.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"use_ocr": "false"},
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["filename"] == "test.docx"
    assert "Test Document Title" in data["text"]
    assert "Body paragraph" in data["text"]
