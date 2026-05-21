"""
DOCX text extraction using python-docx.
"""
from __future__ import annotations

import io
import logging
import os

from schemas import DocumentResult

logger = logging.getLogger(__name__)


class DocxProcessor:
    """Extract text from .docx files using python-docx."""

    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = upload_dir
        os.makedirs(upload_dir, exist_ok=True)

    async def extract_text(
        self, filename: str, content: bytes
    ) -> DocumentResult:
        """Extract text from a DOCX file (bytes) and return a DocumentResult."""
        from docx import Document as DocxDocument

        safe_name = os.path.basename(filename)
        filepath = os.path.join(self.upload_dir, safe_name)
        real_upload = os.path.realpath(self.upload_dir)
        real_file = os.path.realpath(filepath)
        if not real_file.startswith(real_upload + os.sep) and real_file != real_upload:
            raise ValueError("Path traversal detected")
        with open(filepath, "wb") as f:
            f.write(content)

        doc = DocxDocument(io.BytesIO(content))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        all_lines = paragraphs + table_texts
        full_text = "\n".join(all_lines)

        # Count sections as a proxy for "pages"
        section_count = len(doc.sections) if doc.sections else 1

        logger.info(
            "docx extracted: %s (%d paragraphs, %d tables)",
            filename, len(paragraphs), len(doc.tables),
        )

        return DocumentResult(
            filename=filename,
            text=full_text,
            page_count=section_count,
        )
